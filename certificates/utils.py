import uuid
import os
from pathlib import Path
from datetime import datetime
from django.conf import settings
from docx import Document
from docxtpl import DocxTemplate, InlineImage
from django.utils.safestring import mark_safe
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx.shared import Mm

# ---------------- TEMPLATE ----------------
TEMPLATE_MAP = {
    "clearance": "clearance.docx",
    "residency": "residency.docx",
    "indigency": "Indigency.docx",
}

def get_template_path(document_type):
    """
    Return absolute path to template.
    Looks in MEDIA first (uploaded), then default in app templates.
    """
    filename = TEMPLATE_MAP.get(document_type)
    if not filename:
        return None
    uploaded_path = Path(settings.MEDIA_ROOT) / "certificate_templates" / filename
    if uploaded_path.exists():
        return uploaded_path
    default_path = Path(settings.BASE_DIR) / "certificates" / "templates" / "certificate_templates" / filename
    return default_path

def _ensure_dirs():
    MEDIA_ROOT = Path(settings.MEDIA_ROOT)
    (MEDIA_ROOT / "generated" / "docx").mkdir(parents=True, exist_ok=True)
    (MEDIA_ROOT / "generated" / "pdf").mkdir(parents=True, exist_ok=True)
    (MEDIA_ROOT / "qrcodes").mkdir(parents=True, exist_ok=True)
    (MEDIA_ROOT / "certificate_templates").mkdir(parents=True, exist_ok=True)

def docx_to_html(docx_path):
    doc = Document(docx_path)
    html = "".join([f"<p>{para.text}</p>" for para in doc.paragraphs])
    return mark_safe(html)

# ---------------- CERTIFICATE FILES ----------------
def generate_certificate_files(cert):
    """
    Generates DOCX and PDF for a certificate.
    """
    base_dir = Path(settings.MEDIA_ROOT) / "certificates"
    base_dir.mkdir(parents=True, exist_ok=True)

    docx_path = base_dir / f"{cert.unique_id}.docx"
    pdf_path = base_dir / f"{cert.unique_id}.pdf"

    # --- Dates ---
    issued_date = cert.created_at.strftime('%B %d, %Y') if cert.created_at else "N/A"
    reissued_date = cert.reissue_date.strftime('%B %d, %Y') if getattr(cert, 'reissue_date', None) else None

    # --- DOCX ---
    doc = Document()
    doc.add_heading("Barangay Certificate", level=1)
    doc.add_paragraph(f"Name: {cert.full_name or 'N/A'}")
    doc.add_paragraph(f"Address: {cert.address or 'N/A'}")
    doc.add_paragraph(f"Resident Since: {getattr(cert, 'resident_since', 'N/A')}")
    doc.add_paragraph(f"Purpose: {getattr(cert, 'purpose', 'N/A')}")
    doc.add_paragraph(f"Issued Date: {issued_date}")
    if reissued_date:
        doc.add_paragraph(f"Reissued Date: {reissued_date}")
    doc.save(docx_path)

    # --- PDF ---
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    y = 750
    c.setFont("Helvetica-Bold", 16)
    c.drawString(200, y, "Barangay Certificate")
    y -= 40
    c.setFont("Helvetica", 12)
    c.drawString(100, y, f"Name: {cert.full_name or 'N/A'}")
    y -= 20
    c.drawString(100, y, f"Address: {cert.address or 'N/A'}")
    y -= 20
    c.drawString(100, y, f"Resident Since: {getattr(cert, 'resident_since', 'N/A')}")
    y -= 20
    c.drawString(100, y, f"Purpose: {getattr(cert, 'purpose', 'N/A')}")
    y -= 20
    c.drawString(100, y, f"Issued Date: {issued_date}")
    y -= 20
    if reissued_date:
        c.drawString(100, y, f"Reissued Date: {reissued_date}")
    c.showPage()
    c.save()

    # --- Save paths to model ---
    cert.generated_docx.name = f"certificates/{cert.unique_id}.docx"
    cert.generated_pdf.name = f"certificates/{cert.unique_id}.pdf"
    cert.save(update_fields=["generated_docx", "generated_pdf"])
    return True

def generate_unique_id(document_type, reissue_count=0):
    """
    Generate unique ID for certificates.
    """
    prefix_map = {
        "clearance": "CLR",
        "residency": "RES",
        "indigency": "IND",
    }
    prefix = prefix_map.get(document_type, "DOC")
    date_part = datetime.now().strftime("%Y%m%d")
    unique_part = uuid.uuid4().hex[:6].upper()
    if reissue_count > 0:
        return f"{prefix}-{date_part}-{unique_part}-R{reissue_count}"
    return f"{prefix}-{date_part}-{unique_part}"

# certificates/utils.py

def convert(input_path, output_path=None):
    if output_path:
        # use docx2pdf to convert to output_path
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(input_path, output_path)
    else:
        # default behavior
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(input_path)


# certificates/utils.py

def save_data(data):
    """
    Temporary placeholder for saving certificate data.
    Replace this later with your actual form handling logic.
    """
    print("save_data() called with:", data)
    # You can perform validation or database save here later.
    return True
